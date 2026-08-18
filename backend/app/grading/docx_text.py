import io

import docx


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Pull readable text out of a .docx file, in document order.

    Answer keys are sometimes laid out as tables (question | answer | points)
    rather than plain paragraphs, so both are walked and interleaved via
    `document.element.body` rather than reading `document.paragraphs` and
    `document.tables` separately, which would lose the original ordering.
    """
    document = docx.Document(io.BytesIO(file_bytes))

    lines: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = docx.text.paragraph.Paragraph(child, document)
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        elif child.tag.endswith("}tbl"):
            table = docx.table.Table(child, document)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))

    return "\n".join(lines)
